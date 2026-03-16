#
#  Copyright 2025 The InfiniFlow Authors. All Rights Reserved.
#
#  Licensed under the Apache License, Version 2.0 (the "License");
#  you may not use this file except in compliance with the License.
#  You may obtain a copy of the License at
#
#      http://www.apache.org/licenses/LICENSE-2.0
#
#  Unless required by applicable law or agreed to in writing, software
#  distributed under the License is distributed on an "AS IS" BASIS,
#  WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#  See the License for the specific language governing permissions and
#  limitations under the License.
#
"""Word document (.docx) parser for the RAG pipeline.

Extracts text paragraphs and table content from DOCX files. Tables are
converted into a linearized text format with header-value pairs for
downstream chunking and embedding.
"""

from docx import Document
import re
import pandas as pd
from collections import Counter
from rag.nlp import rag_tokenizer
from io import BytesIO


class RAGFlowDocxParser:
    """Parser for Microsoft Word (.docx) documents.

    Extracts paragraphs (with page break tracking) and tables from DOCX files.
    Tables are converted to a text format where each row is represented as
    semicolon-separated "header: value" pairs.
    """

    def __extract_table_content(self, tb):
        """Extract and format content from a DOCX table element.

        Args:
            tb: A python-docx Table object.

        Returns:
            A list of formatted text strings representing the table content.
        """
        df = []
        for row in tb.rows:
            df.append([c.text for c in row.cells])
        return self.__compose_table_content(pd.DataFrame(df))

    def __compose_table_content(self, df):
        """Convert a DataFrame table into linearized text with header-value pairs.

        Analyzes the table structure to identify header rows and data rows,
        then formats each data row as "header1: value1; header2: value2; ...".
        Uses block type classification (date, number, text, etc.) to determine
        which rows are headers vs data.

        Args:
            df: A pandas DataFrame representing the table.

        Returns:
            A list of text strings, each representing one or more table rows.
        """

        def blockType(b):
            """Classify a cell value into a block type for table analysis.

            Returns a type code: Dt (date), Nu (number), Ca (catalog),
            En (english), NE (number+english), Sg (single char), Tx (text),
            Lx (long text), Nr (person name), Ot (other).
            """
            pattern = [
                ("^(20|19)[0-9]{2}[年/-][0-9]{1,2}[月/-][0-9]{1,2}日*$", "Dt"),
                (r"^(20|19)[0-9]{2}年$", "Dt"),
                (r"^(20|19)[0-9]{2}[年/-][0-9]{1,2}月*$", "Dt"),
                ("^[0-9]{1,2}[月/-][0-9]{1,2}日*$", "Dt"),
                (r"^第*[一二三四1-4]季度$", "Dt"),
                (r"^(20|19)[0-9]{2}年*[一二三四1-4]季度$", "Dt"),
                (r"^(20|19)[0-9]{2}[ABCDE]$", "DT"),
                ("^[0-9.,+%/ -]+$", "Nu"),
                (r"^[0-9A-Z/\._~-]+$", "Ca"),
                (r"^[A-Z]*[a-z' -]+$", "En"),
                (r"^[0-9.,+-]+[0-9A-Za-z/$￥%<>（）()' -]+$", "NE"),
                (r"^.{1}$", "Sg")
            ]
            for p, n in pattern:
                if re.search(p, b):
                    return n
            tks = [t for t in rag_tokenizer.tokenize(b).split() if len(t) > 1]
            if len(tks) > 3:
                if len(tks) < 12:
                    return "Tx"
                else:
                    return "Lx"

            if len(tks) == 1 and rag_tokenizer.tag(tks[0]) == "nr":
                return "Nr"

            return "Ot"

        if len(df) < 2:
            return []
        # Determine the dominant cell type in the table (excluding header)
        max_type = Counter([blockType(str(df.iloc[i, j])) for i in range(
            1, len(df)) for j in range(len(df.iloc[i, :]))])
        max_type = max(max_type.items(), key=lambda x: x[1])[0]

        colnm = len(df.iloc[0, :])
        hdrows = [0]  # header is not necessarily appear in the first line
        # If dominant type is numeric, find additional header rows
        if max_type == "Nu":
            for r in range(1, len(df)):
                tys = Counter([blockType(str(df.iloc[r, j]))
                              for j in range(len(df.iloc[r, :]))])
                tys = max(tys.items(), key=lambda x: x[1])[0]
                if tys != max_type:
                    hdrows.append(r)

        # Build formatted output lines with header-value pairs
        lines = []
        for i in range(1, len(df)):
            if i in hdrows:
                continue
            # Find the closest header row(s) above this data row
            hr = [r - i for r in hdrows]
            hr = [r for r in hr if r < 0]
            t = len(hr) - 1
            while t > 0:
                if hr[t] - hr[t - 1] > 1:
                    hr = hr[t:]
                    break
                t -= 1
            # Build header text for each column
            headers = []
            for j in range(len(df.iloc[i, :])):
                t = []
                for h in hr:
                    x = str(df.iloc[i + h, j]).strip()
                    if x in t:
                        continue
                    t.append(x)
                t = ",".join(t)
                if t:
                    t += ": "
                headers.append(t)
            # Combine header and value for each non-empty cell
            cells = []
            for j in range(len(df.iloc[i, :])):
                if not str(df.iloc[i, j]):
                    continue
                cells.append(headers[j] + str(df.iloc[i, j]))
            lines.append(";".join(cells))

        # For wide tables (>3 cols), return each row separately;
        # for narrow tables, join all rows into a single string
        if colnm > 3:
            return lines
        return ["\n".join(lines)]

    def __call__(self, fnm, from_page=0, to_page=100000000):
        """Parse a DOCX file and extract paragraphs and tables.

        Tracks page breaks via Word's lastRenderedPageBreak XML elements
        to support page-range filtering.

        Args:
            fnm: File path (string) or binary content (bytes) of the .docx file.
            from_page: Starting page number (0-based, inclusive).
            to_page: Ending page number (0-based, exclusive).

        Returns:
            A tuple of (sections, tables):
            - sections: list of (text, style_name) tuples for each paragraph
            - tables: list of formatted table content lists
        """
        self.doc = Document(fnm) if isinstance(
            fnm, str) else Document(BytesIO(fnm))
        pn = 0 # parsed page
        secs = [] # parsed contents
        for p in self.doc.paragraphs:
            if pn > to_page:
                break

            runs_within_single_paragraph = [] # save runs within the range of pages
            for run in p.runs:
                if pn > to_page:
                    break
                if from_page <= pn < to_page and p.text.strip():
                    runs_within_single_paragraph.append(run.text) # append run.text first

                # wrap page break checker into a static method
                if 'lastRenderedPageBreak' in run._element.xml:
                    pn += 1

            secs.append(("".join(runs_within_single_paragraph), p.style.name if hasattr(p.style, 'name') else '')) # then concat run.text as part of the paragraph

        tbls = [self.__extract_table_content(tb) for tb in self.doc.tables]
        return secs, tbls
